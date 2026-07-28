#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap
from docx.shared import Inches

SRC_MD = os.path.join(os.path.dirname(__file__), '..', 'doc', 'research_report_final.md')
OUT_DIR = os.path.expanduser('~/Desktop/工作站')
OUT_PATH = os.path.join(OUT_DIR, 'research_report_final.docx')
FORMULA_DIR = os.path.join(os.path.dirname(__file__), '..', 'doc', 'formula_images')

os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(FORMULA_DIR, exist_ok=True)

def set_run_font(run, name='Times New Roman', size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_caption_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    set_run_font(r, name='Times New Roman', size=10, italic=True)
    return p


def render_formula_image(formula: str) -> str:
    try:
        import matplotlib.pyplot as plt
    except ImportError:
        raise ImportError('matplotlib is required to render formulas')
    formula = formula.strip()
    output_path = os.path.join(FORMULA_DIR, f'formula_{abs(hash(formula))}.png')
    if not os.path.exists(output_path):
        fig = plt.figure(figsize=(len(formula) * 0.12 + 0.2, 0.6))
        fig.text(0, 0, f'${formula}$', fontsize=14)
        plt.axis('off')
        fig.savefig(output_path, bbox_inches='tight', pad_inches=0.05, dpi=300)
        plt.close(fig)
    return output_path


def add_code_textbox(doc, code_lines):
    # Create a table with one cell and set it as a textbox-like container
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6)
    cell = table.cell(0, 0)
    # set shading light gray
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F7F7F7')
    tcPr.append(shd)
    for cl in code_lines:
        p = cell.add_paragraph()
        r = p.add_run(cl)
        set_run_font(r, name='Courier New', size=10)
    return table


def main():
    with open(SRC_MD, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()
    # 默认正文字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    in_code = False
    code_lines = []
    in_formula = False
    formula_lines = []

    for line in lines:
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_textbox(doc, code_lines)
                in_code = False
            continue
        if in_code:
            code_lines.append(line)
            continue

        if line.strip() == '$$':
            if not in_formula:
                in_formula = True
                formula_lines = []
                continue
            else:
                formula = ' '.join(fl.strip() for fl in formula_lines)
                img_path = render_formula_image(formula)
                doc.add_picture(img_path, width=Inches(5))
                in_formula = False
                continue
        if in_formula:
            formula_lines.append(line)
            continue

        # 简单处理标题与小节
        if line.strip().endswith('：') and len(line.strip()) < 40 and line.strip().find(' ') == -1:
            # treat as section title
            h = doc.add_heading(level=1)
            r = h.add_run(line.strip())
            set_run_font(r, name='Times New Roman', size=14, bold=True)
            continue
        if line.startswith('图 ') or line.startswith('表 '):
            # 图表注：小两号并斜体
            add_caption_paragraph(doc, line.strip())
            continue
        if line.strip() == '':
            doc.add_paragraph('')
            continue
        # 普通段落
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_run_font(r, name='Times New Roman', size=12)

    doc.save(OUT_PATH)
    print('Saved:', OUT_PATH)

if __name__ == '__main__':
    main()
